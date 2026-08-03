# -*- coding: utf-8 -*-
import re, json, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import content as C

BASE=os.environ.get('JAAI_BASE', os.path.dirname(os.path.abspath(__file__)))
BLIND=os.environ.get('JAAI_BLIND')=='1'      # blinded review copy: omit author block
HIGHLIGHT=os.environ.get('JAAI_HIGHLIGHT')=='1'  # mark <hl>..</hl> spans (revised text) in red
REVCOLOR=RGBColor(0xC0,0x00,0x00)            # revision-marking colour
FIG=BASE+'/fig'; EQ=BASE+'/eq'
SERIF='Times New Roman'

# page geometry (DXA) from template
PGW,PGH=11906,16838
MT,MR,MB,ML=1418,1134,851,1134
CONTENT_DXA=PGW-ML-MR          # 9638
COL_SPACE=425
COL_DXA=(CONTENT_DXA-COL_SPACE)//2   # 4606
CONTENT_IN=CONTENT_DXA/1440.0        # 6.693
COL_IN=COL_DXA/1440.0                # 3.198

doc=Document(BASE+'/submit_template.docx')
body=doc.element.body

# --- clear body but keep final sectPr ---
final_sectPr=body.find(qn('w:sectPr'))
for ch in list(body):
    if ch is final_sectPr: continue
    body.remove(ch)

# ---------- helpers ----------
def set_run(r, font=SERIF, size=None, bold=None, italic=None, sup=False, color=None):
    r.font.name=font
    rPr=r._element.get_or_add_rPr()
    rfonts=rPr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts=OxmlElement('w:rFonts'); rPr.insert(0,rfonts)
    for a in ('w:ascii','w:hAnsi','w:cs'): rfonts.set(qn(a),font)
    # keep east-asian font for any hangul
    rfonts.set(qn('w:eastAsia'),'Batang')
    if size is not None: r.font.size=Pt(size)
    if bold is not None: r.font.bold=bold
    if italic is not None: r.font.italic=italic
    if color is not None: r.font.color.rgb=color
    if sup: r.font.superscript=True

CITE=re.compile(r'\[([\d,\s–-]+)\]')
def add_runs(p, text, size=10, italic_default=False):
    # parse <i>..</i> italics, <hl>..</hl> revision marks, and [n] citations -> superscript (n)
    parts=re.split(r'(<i>|</i>|<hl>|</hl>)', text)
    ital=italic_default; hl=False
    for part in parts:
        if part=='<i>': ital=True; continue
        if part=='</i>': ital=False; continue
        if part=='<hl>': hl=True; continue
        if part=='</hl>': hl=False; continue
        if not part: continue
        col=REVCOLOR if (hl and HIGHLIGHT) else None
        # within part, handle citations
        idx=0
        for m in CITE.finditer(part):
            if m.start()>idx:
                r=p.add_run(part[idx:m.start()]); set_run(r,size=size,italic=ital,color=col)
            r=p.add_run('('+m.group(1).replace(' ','')+')'); set_run(r,size=size,sup=True,color=col)
            idx=m.end()
        if idx<len(part):
            r=p.add_run(part[idx:]); set_run(r,size=size,italic=ital,color=col)

def kill_num(p):
    pPr=p._p.get_or_add_pPr()
    for e in pPr.findall(qn('w:numPr')): pPr.remove(e)
    numPr=OxmlElement('w:numPr')
    ilvl=OxmlElement('w:ilvl'); ilvl.set(qn('w:val'),'0')
    numId=OxmlElement('w:numId'); numId.set(qn('w:val'),'0')
    numPr.append(ilvl); numPr.append(numId)
    pStyle=pPr.find(qn('w:pStyle'))
    if pStyle is not None: pStyle.addnext(numPr)
    else: pPr.insert(0,numPr)

def para(style=None, align=None, space_before=None, space_after=None):
    p=doc.add_paragraph()
    if style:
        try: p.style=doc.styles[style]
        except Exception: pass
    kill_num(p)
    if align is not None: p.alignment=align
    pf=p.paragraph_format
    if space_before is not None: pf.space_before=Pt(space_before)
    if space_after is not None: pf.space_after=Pt(space_after)
    return p

def img_size(path, max_w_in, max_h_in=9.0):
    im=Image.open(path); w,h=im.size; ar=w/h
    W=max_w_in; H=W/ar
    if H>max_h_in: H=max_h_in; W=H*ar
    return W,H

def add_image(path, max_w_in, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=2):
    W,H=img_size(path,max_w_in)
    p=para(align=align, space_before=space_before, space_after=space_after)
    p.paragraph_format.line_spacing=1.0
    r=p.add_run(); r.add_picture(path, width=Inches(W))
    return p

def add_equation(key, maxw):
    add_image(f'{EQ}/{key.replace("eq","e")}.png', min(maxw, maxw), align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=4)

import json as _json
from lxml import etree as _ET
OMML=_json.load(open(BASE+'/omml.json'))
def _parse_omml(s):
    w='<root xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'+s+'</root>'
    return _ET.fromstring(w)[0]
def add_omml_equation(n, full=True):
    width = CONTENT_IN if full else COL_IN
    p=para(space_before=3, space_after=3)
    p.paragraph_format.line_spacing=1.0
    pf=p.paragraph_format
    pf.tab_stops.add_tab_stop(Inches(width/2), WD_TAB_ALIGNMENT.CENTER)
    pf.tab_stops.add_tab_stop(Inches(width-0.05), WD_TAB_ALIGNMENT.RIGHT)
    r1=p.add_run(); r1._r.append(OxmlElement('w:tab'))
    p._p.append(_parse_omml(OMML[str(n)]))
    r2=p.add_run(); r2._r.append(OxmlElement('w:tab'))
    r3=p.add_run('('+str(n)+')'); set_run(r3,size=10)
    return p

def add_caption_fig(label, cap, size=9):
    p=para(align=WD_ALIGN_PARAGRAPH.CENTER, space_before=3, space_after=8)
    p.paragraph_format.line_spacing=1.0
    r=p.add_run(label+'  '); set_run(r,size=size,bold=True)
    add_runs(p, cap, size=size)

def add_caption_tbl(label, cap, size=9):
    p=para(align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=3)
    p.paragraph_format.line_spacing=1.0
    r=p.add_run(label+'  '); set_run(r,size=size,bold=True)
    add_runs(p, cap, size=size)

def _set_cell_border(cell, **kw):
    tcPr=cell._tc.get_or_add_tcPr()
    tb=tcPr.find(qn('w:tcBorders'))
    if tb is None:
        tb=OxmlElement('w:tcBorders'); tcPr.append(tb)
    for edge in ('top','bottom','left','right'):
        if edge in kw:
            e=tb.find(qn('w:'+edge))
            if e is None:
                e=OxmlElement('w:'+edge); tb.append(e)
            sz,val=kw[edge]
            e.set(qn('w:val'),'single' if val else 'nil')
            e.set(qn('w:sz'),str(sz)); e.set(qn('w:space'),'0'); e.set(qn('w:color'),'000000')

def add_table(label, cap, header, rows, span):
    add_caption_tbl(label, cap)
    total_in = CONTENT_IN-0.1 if span else COL_IN-0.05
    ncol=len(header)
    t=doc.add_table(rows=1+len(rows), cols=ncol)
    t.alignment=1  # center
    t.autofit=False
    # column widths
    colw=Emu(int(Inches(total_in)/ncol))
    for row in t.rows:
        for c in row.cells: c.width=colw
    # set grid
    tblPr=t._tbl.tblPr
    tw=OxmlElement('w:tblW'); tw.set(qn('w:w'),str(int(total_in*1440))); tw.set(qn('w:type'),'dxa'); tblPr.append(tw)
    def fill(row_cells, values, bold=False, size=9):
        for c,val in zip(row_cells, values):
            c.text=''
            p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(1)
            p.paragraph_format.line_spacing=1.0
            add_runs(p, str(val), size=size)
            for r in p.runs:
                set_run(r,size=size,bold=bold)
    fill(t.rows[0].cells, header, bold=True)
    for i,rw in enumerate(rows):
        fill(t.rows[1+i].cells, rw)
    # booktabs borders: top of header, bottom of header, bottom of last
    nrow=len(t.rows)
    for j,c in enumerate(t.rows[0].cells):
        _set_cell_border(c, top=(8,True), bottom=(6,True))
    for j,c in enumerate(t.rows[-1].cells):
        _set_cell_border(c, bottom=(8,True))
    # remove internal/vertical borders explicitly (nil) for all cells
    for row in t.rows:
        for c in row.cells:
            _set_cell_border(c, left=(0,False), right=(0,False))
    return t

# ---------- section config ----------
def cfg_section(sec, ncols):
    sp=sec._sectPr
    # type continuous
    for tag in ('w:type',):
        e=sp.find(qn(tag))
        if e is None:
            e=OxmlElement(tag); sp.append(e)
        e.set(qn('w:val'),'continuous')
    # pgSz
    pgsz=sp.find(qn('w:pgSz'))
    if pgsz is None: pgsz=OxmlElement('w:pgSz'); sp.append(pgsz)
    pgsz.set(qn('w:w'),str(PGW)); pgsz.set(qn('w:h'),str(PGH)); pgsz.set(qn('w:code'),'9')
    # pgMar
    m=sp.find(qn('w:pgMar'))
    if m is None: m=OxmlElement('w:pgMar'); sp.append(m)
    m.set(qn('w:top'),str(MT)); m.set(qn('w:right'),str(MR)); m.set(qn('w:bottom'),str(MB))
    m.set(qn('w:left'),str(ML)); m.set(qn('w:header'),'1134'); m.set(qn('w:footer'),'397'); m.set(qn('w:gutter'),'0')
    # cols
    cols=sp.find(qn('w:cols'))
    if cols is None: cols=OxmlElement('w:cols'); sp.append(cols)
    for a in list(cols.attrib): del cols.attrib[qn(a) if not a.startswith('{') else a]
    if ncols==2:
        cols.set(qn('w:num'),'2'); cols.set(qn('w:space'),str(COL_SPACE)); cols.set(qn('w:equalWidth'),'1')
    else:
        cols.set(qn('w:space'),'720')

def strip_hf_refs(sec):
    sp=sec._sectPr
    for tag in ('w:headerReference','w:footerReference','w:titlePg'):
        for e in sp.findall(qn(tag)): sp.remove(e)

def set_masthead_refs(sec):
    sp=sec._sectPr
    strip_hf_refs(sec)
    def ref(tag,typ,rid):
        e=OxmlElement(tag); e.set(qn('w:type'),typ); e.set(qn('r:id'),rid); sp.insert(0,e)
    # order: header first then footer, titlePg after
    ref('w:headerReference','default','rId8')
    ref('w:headerReference','first','rId11')
    ref('w:footerReference','default','rId10')
    ref('w:footerReference','first','rId12')
    tp=OxmlElement('w:titlePg'); sp.append(tp)

# ================= BUILD CONTENT =================
# --- Title block (section 0, single column) ---
ORANGE=RGBColor(0xDE,0x6B,0x1E)
def author_line(runs, size, ko=False, space_before=0, space_after=2, color=None, bold=False):
    p=para(align=WD_ALIGN_PARAGRAPH.CENTER, space_before=space_before, space_after=space_after)
    p.paragraph_format.line_spacing=1.1
    for text,sup in runs:
        r=p.add_run(text); set_run(r,size=size,sup=sup,bold=bold,color=color,font=('Batang' if ko else SERIF))
    return p

def add_author_ko():
    author_line([('주영환',False),('1,2,*',True)], 11.5, ko=True, space_before=2, space_after=1)
    author_line([('1',True),('한국에너지기술연구원 에너지효율연구본부, ',False),
                 ('2',True),('과학기술연합대학원대학교 에너지공학',False)], 9, ko=True, space_after=4)

def add_author_en():
    author_line([('Younghwan Joo',False),('1,2,*',True)], 11, space_before=2, space_after=1)
    author_line([('1',True),('Energy Efficiency Research Division, Korea Institute of Energy Research, ',False),
                 ('2',True),('Energy Engineering, University of Science and Technology',False)], 9, space_after=1)
    author_line([('*Corresponding author, E-mail: yhjoo@kier.re.kr',False)], 9, color=ORANGE, bold=True, space_before=1, space_after=2)

def title_para(text, size, bold=True, italic=False, ko=False):
    p=para(align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=4)
    p.paragraph_format.line_spacing=1.1
    r=p.add_run(text)
    set_run(r,size=size,bold=bold,italic=italic, font=('Batang' if ko else SERIF))
    return p

para(space_after=2)  # small top gap
title_para(C.KO_TITLE, 15, bold=True, ko=True)
if not BLIND: add_author_ko()
title_para(C.EN_TITLE, 13, bold=True)
if not BLIND: add_author_en()
para(space_after=2)

# --- Abstract box (2-row, 2-col bordered table) ---
box=doc.add_table(rows=2, cols=2); box.autofit=False; box.alignment=1
lw=int(Inches(CONTENT_IN*0.34)); rw=int(Inches(CONTENT_IN*0.64))
tblPr=box._tbl.tblPr
tw=OxmlElement('w:tblW'); tw.set(qn('w:w'),str(lw+rw)); tw.set(qn('w:type'),'dxa'); tblPr.append(tw)
box.rows[0].cells[0].width=Emu(lw); box.rows[0].cells[1].width=Emu(rw)
box.rows[1].cells[0].width=Emu(lw); box.rows[1].cells[1].width=Emu(rw)
# header row
h0=box.rows[0].cells[0]; h0.text=''; p=h0.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('A R T I C L E   I N F O'); set_run(r,size=9,bold=False)
h1=box.rows[0].cells[1]; h1.text=''; p=h1.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('A B S T R A C T'); set_run(r,size=9,bold=False)
# content row - keywords
c0=box.rows[1].cells[0]; c0.text=''
p=c0.paragraphs[0]; p.paragraph_format.space_after=Pt(3)
r=p.add_run('Key Words:'); set_run(r,size=9,italic=True)
p2=c0.add_paragraph(); p2.paragraph_format.space_after=Pt(0); p2.paragraph_format.line_spacing=1.1
r=p2.add_run(C.KEYWORDS); set_run(r,size=9)
# content row - abstract
c1=box.rows[1].cells[1]; c1.text=''
p=c1.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing=1.15
add_runs(p, C.ABSTRACT, size=9)
# box borders: top of row0, bottom of row0 (header underline), bottom of row1
for c in box.rows[0].cells: _set_cell_border(c, top=(8,True), bottom=(6,True), left=(0,False), right=(0,False))
for c in box.rows[1].cells: _set_cell_border(c, bottom=(8,True), left=(0,False), right=(0,False), top=(0,False))
# cell margins
for row in box.rows:
    for c in row.cells:
        tcPr=c._tc.get_or_add_tcPr(); mar=OxmlElement('w:tcMar')
        for edge,val in (('top',60),('bottom',60),('left',120),('right',120)):
            e=OxmlElement('w:'+edge); e.set(qn('w:w'),str(val)); e.set(qn('w:type'),'dxa'); mar.append(e)
        tcPr.append(mar)
para(space_after=2)

# ---- end section 0 (title/abstract, 1 col) ----
STYLE={'sec':'a0','secx':'a0','secstar':'af7','sub':'af7','ssub':'a2'}

def emit_inline(b):
    """emit a block that stays within the current (2-col) flow"""
    kind=b[0]
    if kind in ('sec','secx','secstar'):
        if kind=='sec':
            num,title=b[1],b[2]; txt=f'{num}. {title}'; align=WD_ALIGN_PARAGRAPH.CENTER
        elif kind=='secx':
            txt=b[1]; align=WD_ALIGN_PARAGRAPH.CENTER
        else:
            txt=b[1]; align=WD_ALIGN_PARAGRAPH.LEFT
        p=para(style=STYLE[kind], align=align, space_before=10, space_after=5)
        r=p.add_run(txt); set_run(r,size=11 if kind!='secstar' else 10,bold=True)
    elif kind=='sub':
        p=para(style='af7', align=WD_ALIGN_PARAGRAPH.LEFT, space_before=8, space_after=4)
        r=p.add_run(f'{b[1]} {b[2]}'); set_run(r,size=10.5,bold=True)
    elif kind=='ssub':
        p=para(style='a2', align=WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=3)
        r=p.add_run(f'{b[1]} {b[2]}'); set_run(r,size=10,bold=True,italic=True)
    elif kind=='p':
        p=para(style='a1', align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
        p.paragraph_format.line_spacing=1.15
        p.paragraph_format.first_line_indent=Inches(0.16)
        add_runs(p, b[1], size=10)
    elif kind=='eq':
        add_equation(b[1], COL_IN-0.05)
    elif kind=='fig':   # inline (1-col) figure
        add_image(f'{FIG}/{b[1]}.png', COL_IN-0.08)
        add_caption_fig(b[2], b[3])
    elif kind=='tbl':   # inline (2-col) table
        add_table(b[1], b[2], b[3], b[4], span=False)

def is_full(b):
    if b[0]=='fig' and b[4]: return True
    if b[0]=='tbl' and len(b[3])>=3: return True
    return False

def emit_full(b):
    if b[0]=='fig':
        add_image(f'{FIG}/{b[1]}.png', CONTENT_IN-0.1)
        add_caption_fig(b[2], b[3])
    elif b[0]=='tbl':
        add_table(b[1], b[2], b[3], b[4], span=True)

# Walk blocks, creating sections at column-context changes.
# Current context after title/abstract will become 2-col once we start body.
# We use add_section before switching contexts.
# build flow: merge consecutive equation blocks into one full-width group
# '_nb'-suffixed blocks (e.g. funding acknowledgement) are dropped from the blinded copy
_B=[b for b in C.BLOCKS if not (BLIND and b[0].endswith('_nb'))]
_B=[((b[0][:-3],)+tuple(b[1:]) if b[0].endswith('_nb') else b) for b in _B]
flow=[]; _i=0
while _i<len(_B):
    b=_B[_i]
    if b[0]=='eq':
        grp=[]
        while _i<len(_B) and _B[_i][0]=='eq':
            grp.append(_B[_i][2]); _i+=1
        flow.append(('eqgroup',grp))
    else:
        flow.append(('blk',b)); _i+=1

def item_full(it):
    if it[0]=='eqgroup': return True
    b=it[1]
    if b[0]=='fig' and b[4]: return True
    if b[0]=='tbl' and len(b[3])>=3: return True
    return False

sec=doc.add_section(WD_SECTION.CONTINUOUS)  # first 2-col body section
for it in flow:
    if item_full(it):
        doc.add_section(WD_SECTION.CONTINUOUS)
        if it[0]=='eqgroup':
            for n in it[1]: add_omml_equation(n, full=True)
        else:
            emit_full(it[1])
        doc.add_section(WD_SECTION.CONTINUOUS)
    else:
        emit_inline(it[1])

# ---- References ----
p=para(style='af9', align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=5)
r=p.add_run('References'); set_run(r,size=11,bold=True)
refs=json.load(open(BASE+'/refs.json', encoding='utf-8'))
for i,segs in enumerate(refs,1):
    p=para(style='a4', align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2)
    p.paragraph_format.line_spacing=1.05
    p.paragraph_format.left_indent=Inches(0.16)
    p.paragraph_format.first_line_indent=Inches(-0.16)
    r=p.add_run(f'({i}) '); set_run(r,size=9)
    for txt,ital in segs:
        r=p.add_run(txt); set_run(r,size=9,italic=ital)

# ================= CONFIGURE SECTIONS =================
secs=doc.sections
# secs[0] = title/abstract (1 col, masthead). secs[1..] alternate.
# Determine column count per section by inspecting: section 0 ->1col.
# We know pattern: after title, body sections alternate; sections created:
#  - the doc.add_section before loop -> starts section index1 (2col body)
#  - each full block adds two sections: one 1col(full) + one 2col(after)
# Easiest: set section 0 = 1col; then walk remaining and infer via a rebuilt list.
# Instead track counts precisely:
col_plan=[1,2]  # title, first body
for it in flow:
    if item_full(it):
        col_plan.append(1); col_plan.append(2)
assert len(col_plan)==len(secs), f'plan {len(col_plan)} != secs {len(secs)}'
for i,(sec,nc) in enumerate(zip(secs, col_plan)):
    cfg_section(sec, nc)
    if i==0:
        set_masthead_refs(sec)
    else:
        strip_hf_refs(sec)

doc.save(BASE+'/build1.docx')
print('sections:', len(secs), 'plan head:', col_plan[:8])
print('saved build1.docx')
